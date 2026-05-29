# -*- coding: utf-8 -*-
"""Update Part 2 Spy SQL answers in Homework 4 docx to match ERD (homework4 image - spy.png)."""
from pathlib import Path

from docx import Document
from docx.shared import Pt

REPO_ROOT = Path(__file__).resolve().parents[2]
DOC = REPO_ROOT / "Homework" / "docx" / "CS 486_586 – Homework 4-emmanart.docx"

# Lines to replace block starting at "Answer 6:" through before "Part 3:"
SPY_ANSWER_LINES = [
    "Answer 6:",
    "SQL (subquery) — table team (team_id, name) from ERD:",
    "SELECT name",
    "FROM team",
    "WHERE team_id = (",
    "  SELECT MAX(team_id)",
    "  FROM team",
    ");",
    "",
    "Rows returned: [paste count after you run]",
    "[SCREENSHOT: first 5 rows of query result]",
    "",
    "Answer 7:",
    "SQL (subquery) — path from ERD: agent -> teamrel -> team <- mission; mission.access_id -> securityclearance:",
    "SELECT DISTINCT a.first, a.last",
    "FROM agent AS a",
    "WHERE a.agent_id IN (",
    "  SELECT tr.agent_id",
    "  FROM teamrel AS tr",
    "  WHERE tr.team_id IN (",
    "    SELECT m.team_id",
    "    FROM mission AS m",
    "    INNER JOIN securityclearance AS sc",
    "      ON m.access_id = sc.sc_id",
    "    WHERE sc.sc_level = 'Top Secret'",
    "      AND lower(m.mission_status) = 'failed'",
    "  )",
    ");",
    "",
    "Rows returned: [paste count after you run]",
    "[SCREENSHOT: first 5 rows — or all rows if fewer than 5]",
    "",
    "Answer 8:",
    "SQL (subquery) — tables language and languagerel (lang_id, agent_id) from ERD:",
    "SELECT l.language, COUNT(*) AS number_of_speakers",
    "FROM language AS l",
    "INNER JOIN languagerel AS lr ON l.lang_id = lr.lang_id",
    "GROUP BY l.lang_id, l.language",
    "HAVING COUNT(*) = (",
    "  SELECT MAX(speaker_count)",
    "  FROM (",
    "    SELECT COUNT(*) AS speaker_count",
    "    FROM languagerel",
    "    GROUP BY lang_id",
    "  ) AS counts_per_language",
    ");",
    "",
    "Rows returned: [paste count after you run]",
    "[SCREENSHOT: first 5 rows — or all rows if fewer than 5]",
    "",
    "Explanation: Each row in languagerel links one agent_id to one lang_id; COUNT is number of agents per language.",
]


def set_paragraph(paragraph, text, *, mono=False):
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = text
    run.font.name = "Consolas" if mono else "Calibri"
    run.font.size = Pt(9 if mono else 11)


def is_sql_line(line: str) -> bool:
    s = line.strip()
    if not s:
        return False
    if s.startswith("SQL") or s.startswith("Explanation") or s.startswith("Rows") or s.startswith("["):
        return False
    return (
        s.startswith("SELECT")
        or s.startswith("FROM")
        or s.startswith("WHERE")
        or s.startswith("INNER")
        or s.startswith("GROUP")
        or s.startswith("HAVING")
        or s.startswith(");")
        or s.startswith("  ")
        or s == ");"
    )


def main():
    doc = Document(DOC)
    paras = doc.paragraphs

    start = end = None
    for i, p in enumerate(paras):
        t = p.text.strip()
        if t == "Answer 6:":
            start = i
        if start is not None and t.startswith("Part 3:"):
            end = i
            break

    if start is None or end is None:
        raise RuntimeError("Could not find Answer 6 / Part 3 block in docx")

    needed = len(SPY_ANSWER_LINES)
    available = end - start
    if needed > available:
        raise RuntimeError(f"Need {needed} paragraphs but only {available} slots ({start}-{end})")

    for j, line in enumerate(SPY_ANSWER_LINES):
        set_paragraph(paras[start + j], line, mono=is_sql_line(line))

    for j in range(start + needed, end):
        set_paragraph(paras[j], "")

    doc.save(DOC)
    print(f"Updated Spy queries in {DOC} (paragraphs {start}-{start + needed - 1})")


if __name__ == "__main__":
    main()
