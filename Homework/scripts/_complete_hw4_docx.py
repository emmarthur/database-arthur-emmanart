# -*- coding: utf-8 -*-
"""Fill CS 486/586 Homework 4 docx with answers and screenshot placeholders."""
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

REPO_ROOT = Path(__file__).resolve().parents[2]
DOC_PATH = REPO_ROOT / "Homework" / "CS 486_586 – Homework 4-emmanart.docx"
HEADING_BLACK = RGBColor(0, 0, 0)


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = paragraph._element.makeelement(qn("w:p"), {})
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        run = new_para.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    return new_para


def add_block_after(paragraph, lines, *, sql_blocks=None):
    """Insert lines after paragraph; sql_blocks is list of (start_idx, end_idx) inclusive in lines."""
    sql_blocks = sql_blocks or []
    sql_set = set()
    for start, end in sql_blocks:
        sql_set.update(range(start, end + 1))
    current = paragraph
    for i, line in enumerate(lines):
        current = insert_paragraph_after(current, "")
        if i in sql_set:
            run = current.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            run = current.add_run(line)
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            if line.startswith("Answer") or line.startswith("Question"):
                run.bold = True
    return current


def force_heading_black(doc):
    for p in doc.paragraphs:
        if p.style and p.style.name and p.style.name.startswith("Heading"):
            for r in p.runs:
                r.font.color.rgb = HEADING_BLACK


def main():
    doc = Document(DOC_PATH)

    # Solo student note
    doc.paragraphs[2].text = "Student 2 name and ODIN id: N/A (individual submission)"

    # Blocks inserted bottom-to-top so paragraph indices stay valid
    insertions = []

    insertions.append(
        (
            5,
            [
                "",
                "Contribution table:",
                "Emmanuel Arthur (emmanart) — completed all questions in Parts 1–4.",
                "",
            ],
            None,
        )
    )

    q1 = [
        "",
        "Answer 1:",
        "SQL (subquery):",
        "SELECT Title",
        "FROM Books",
        "WHERE PubID IN (",
        "  SELECT ID",
        "  FROM Publishers",
        "  WHERE Name = 'Viking'",
        ");",
        "",
        "Explanation: The inner query finds the publisher id for Viking. The outer query returns book titles for that publisher.",
        "",
    ]
    insertions.append((17, q1, [(3, 9)]))

    q2 = [
        "",
        "Answer 2:",
        "SQL (subquery):",
        "SELECT DISTINCT Name",
        "FROM Authors",
        "WHERE ID IN (",
        "  SELECT AuthorID",
        "  FROM Books",
        "  WHERE Title LIKE '%COVID%'",
        ");",
        "",
        "Explanation: The inner query finds author ids for books whose title contains COVID. The outer query returns those authors' names.",
        "",
    ]
    insertions.append((18, q2, [(3, 10)]))

    q3 = [
        "",
        "Answer 3:",
        "Question 1 rewritten with JOIN:",
        "SELECT b.Title",
        "FROM Books b",
        "INNER JOIN Publishers p ON b.PubID = p.ID",
        "WHERE p.Name = 'Viking';",
        "",
        "Question 2 rewritten with JOIN:",
        "SELECT DISTINCT a.Name",
        "FROM Authors a",
        "INNER JOIN Books b ON a.ID = b.AuthorID",
        "WHERE b.Title LIKE '%COVID%';",
        "",
        "Explanation: Same results as Questions 1 and 2, but the filter is expressed with joins instead of IN subqueries.",
        "",
    ]
    insertions.append((19, q3, [(3, 6), (8, 12)]))

    q4 = [
        "",
        "Answer 4:",
        "SQL (subquery):",
        "SELECT Title",
        "FROM Books",
        "WHERE PageCount <= (",
        "  SELECT AVG(PageCount) / 2.0",
        "  FROM Books",
        ");",
        "",
        "Explanation: The subquery computes half the average page count. The outer query keeps books at or below that value.",
        "",
    ]
    insertions.append((20, q4, [(3, 9)]))

    q5 = [
        "",
        "Answer 5:",
        "SQL (correlated subquery):",
        "SELECT Genre, Title, PageCount",
        "FROM Books b",
        "WHERE PageCount = (",
        "  SELECT MIN(PageCount)",
        "  FROM Books b2",
        "  WHERE b2.Genre = b.Genre",
        ");",
        "",
        "Explanation: For each book, the subquery finds the minimum page count in the same genre. Matching rows are the shortest book(s); ties are all included.",
        "",
        "Note (Part 1): Create/load a Books database with the schema shown above if you do not already have one. The homework says you may assume Viking and COVID example data exist.",
        "",
    ]
    insertions.append((24, q5, [(3, 10)]))

    spy_intro = [
        "",
        "Part 2 answers (Spy database). I ran each query on my local Spy database from Activity 04-2.",
        "",
    ]
    insertions.append((32, spy_intro, None))

    spy6 = [
        "",
        "Answer 6:",
        "SQL (subquery) — table team (team_id, name) from ERD:",
        "SELECT name",
        "FROM team",
        "WHERE team_id = (",
        "  SELECT MAX(team_id)",
        "  FROM team",
        ");",
        "",
        "Rows returned: [paste count after you run]",
        "[SCREENSHOT: first 5 rows of query result]",
        "",
    ]
    insertions.append((33, spy6, [(3, 9)]))

    spy7 = [
        "",
        "Answer 7:",
        "SQL (subquery) — path from ERD: agent -> teamrel -> team <- mission; mission.access_id -> securityclearance:",
        "SELECT DISTINCT a.first, a.last",
        "FROM agent AS a",
        "WHERE a.agent_id IN (",
        "  SELECT tr.agent_id",
        "  FROM teamrel AS tr",
        "  WHERE tr.team_id IN (",
        "    SELECT m.team_id",
        "    FROM mission AS m",
        "    INNER JOIN securityclearance AS sc",
        "      ON m.access_id = sc.sc_id",
        "    WHERE sc.sc_level = 'Top Secret'",
        "      AND lower(m.mission_status) = 'failed'",
        "  )",
        ");",
        "",
        "Rows returned: [paste count after you run]",
        "[SCREENSHOT: first 5 rows — or all rows if fewer than 5]",
        "",
    ]
    insertions.append((34, spy7, [(4, 20)]))

    spy8 = [
        "",
        "Answer 8:",
        "SQL (subquery) — tables language and languagerel (lang_id, agent_id) from ERD:",
        "SELECT l.language, COUNT(*) AS number_of_speakers",
        "FROM language AS l",
        "INNER JOIN languagerel AS lr ON l.lang_id = lr.lang_id",
        "GROUP BY l.lang_id, l.language",
        "HAVING COUNT(*) = (",
        "  SELECT MAX(speaker_count)",
        "  FROM (",
        "    SELECT COUNT(*) AS speaker_count",
        "    FROM languagerel",
        "    GROUP BY lang_id",
        "  ) AS counts_per_language",
        ");",
        "",
        "Rows returned: [paste count after you run]",
        "[SCREENSHOT: first 5 rows — or all rows if fewer than 5]",
        "",
        "Explanation: Each row in languagerel links one agent_id to one lang_id; COUNT is number of agents per language.",
        "",
    ]
    insertions.append((35, spy8, [(3, 15)]))

    p3_9a = [
        "",
        "Answer 9a:",
        "The statement creates a new table named newagent with the same column names and data types as agent, but with zero rows.",
        "",
    ]
    insertions.append((45, p3_9a, None))

    p3_9b = [
        "",
        "Answer 9b:",
        "WHERE TRUE = FALSE is always false, so no rows from agent are copied. I get the column structure only (CREATE TABLE ... AS SELECT with an empty result).",
        "",
    ]
    insertions.append((46, p3_9b, None))

    p3_9c = [
        "",
        "Answer 9c:",
        "[SCREENSHOT: \\d agent and \\d newagent in psql, or DBeaver column list for both tables showing matching columns]",
        "",
    ]
    insertions.append((47, p3_9c, None))

    p3_9d = [
        "",
        "Answer 9d:",
        "No. CREATE TABLE ... AS SELECT copies columns and data types (and data if any), but it does not copy primary keys, foreign keys, unique constraints, indexes, or defaults from the source table. I would need ALTER TABLE statements to add those constraints to newagent if I wanted them.",
        "",
    ]
    insertions.append((48, p3_9d, None))

    p3_9e = [
        "",
        "Answer 9e:",
        "DROP TABLE newagent;",
        "",
        "[SCREENSHOT optional: message showing table dropped]",
        "",
    ]
    insertions.append((49, p3_9e, [(2, 2)]))

    p4_10 = [
        "",
        "Answer 10:",
        "A database transaction is a group of SQL operations that the database treats as one unit: either all of them take effect together, or none of them do if something goes wrong.",
        "",
    ]
    insertions.append((54, p4_10, None))

    p4_11 = [
        "",
        "Answer 11:",
        "COMMIT saves all changes made during the transaction so they become permanent. ROLLBACK undoes those changes and returns the database to the state before the transaction started.",
        "",
    ]
    insertions.append((55, p4_11, None))

    p4_12 = [
        "",
        "Answer 12:",
        "A transfer should subtract $100 from A and add $100 to B. If those are two separate auto-commit statements, the first could succeed and the second could fail (or vice versa), leaving incorrect balances. A transaction makes sure both updates succeed or neither does.",
        "",
    ]
    insertions.append((56, p4_12, None))

    p4_13 = [
        "",
        "Answer 13:",
        "BEGIN;",
        "UPDATE accounts SET balance = balance - 100 WHERE account_id = 'A';",
        "UPDATE accounts SET balance = balance + 100 WHERE account_id = 'B';",
        "COMMIT;",
        "",
    ]
    insertions.append((57, p4_13, [(2, 5)]))

    p4_14 = [
        "",
        "Answer 14:",
        "If the first UPDATE succeeds but the second fails, the transaction should not be left half-finished. I should ROLLBACK so the deduction from Account A is undone and both accounts stay consistent.",
        "",
    ]
    insertions.append((64, p4_14, None))

    p4_15 = [
        "",
        "Answer 15:",
        "Atomicity: The transfer is all-or-nothing. Either both the withdrawal and deposit happen, or ROLLBACK ensures neither does.",
        "",
        "Consistency: The database moves from one valid state to another (for example total money unchanged). If a constraint like balance >= 0 would be violated, the transaction can ROLLBACK.",
        "",
    ]
    insertions.append((65, p4_15, None))

    for idx, lines, sql_blocks in sorted(insertions, key=lambda x: x[0], reverse=True):
        add_block_after(doc.paragraphs[idx], lines, sql_blocks=sql_blocks)

    force_heading_black(doc)
    doc.save(DOC_PATH)
    print(f"Updated {DOC_PATH}")


if __name__ == "__main__":
    main()
