# -*- coding: utf-8 -*-
"""Qualify Spy homework SQL with schema spy (spy.table)."""
from pathlib import Path

from docx import Document
from docx.shared import Pt

REPO_ROOT = Path(__file__).resolve().parents[2]
DOC = REPO_ROOT / "Homework" / "CS 486_586 – Homework 4-emmanart.docx"

SPY_ANSWER_LINES = [
    "Answer 6:",
    "SQL (subquery) — spy.team (team_id, name) from ERD:",
    "SELECT name",
    "FROM spy.team",
    "WHERE team_id = (",
    "  SELECT MAX(team_id)",
    "  FROM spy.team",
    ");",
    "",
    "Rows returned: [paste count after you run]",
    "[SCREENSHOT: first 5 rows of query result]",
    "",
    "Answer 7:",
    "SQL (subquery) — spy schema from ERD: spy.agent -> spy.teamrel -> spy.team <- spy.mission; spy.mission.access_id -> spy.securityclearance:",
    "SELECT DISTINCT a.first, a.last",
    "FROM spy.agent AS a",
    "WHERE a.agent_id IN (",
    "  SELECT tr.agent_id",
    "  FROM spy.teamrel AS tr",
    "  WHERE tr.team_id IN (",
    "    SELECT m.team_id",
    "    FROM spy.mission AS m",
    "    INNER JOIN spy.securityclearance AS sc",
    "      ON m.access_id = sc.sc_id",
    "    WHERE sc.sc_level = 'Top Secret'",
    "      AND lower(m.mission_status) = 'failed'",
    "  )",
    ");",
    "",
    "Rows returned: [paste count after you run]",
    "[SCREENSHOT: first 5 rows — or all rows if fewer than 5]",
    "",
    "Answer 8:",
    "SQL (subquery) — spy.language and spy.languagerel (lang_id, agent_id) from ERD:",
    "SELECT l.language, COUNT(*) AS number_of_speakers",
    "FROM spy.language AS l",
    "INNER JOIN spy.languagerel AS lr ON l.lang_id = lr.lang_id",
    "GROUP BY l.lang_id, l.language",
    "HAVING COUNT(*) = (",
    "  SELECT MAX(speaker_count)",
    "  FROM (",
    "    SELECT COUNT(*) AS speaker_count",
    "    FROM spy.languagerel",
    "    GROUP BY lang_id",
    "  ) AS counts_per_language",
    ");",
    "",
    "Rows returned: [paste count after you run]",
    "[SCREENSHOT: first 5 rows — or all rows if fewer than 5]",
    "",
    "Explanation: Each row in spy.languagerel links one agent_id to one lang_id; COUNT is number of agents per language.",
]

Q9_REPLACEMENTS = {
    "FROM agent": "FROM spy.agent",
    "CREATE TABLE newagent AS": "CREATE TABLE spy.newagent AS",
    "DROP TABLE newagent;": "DROP TABLE spy.newagent;",
    "The statement creates a new table named newagent with the same column names and data types as agent, but with zero rows. From the Spy ERD, agent has: agent_id, first, middle, last, address, city, country, salary, clearance_id — newagent gets those nine columns with matching types and no data.": (
        "The statement creates spy.newagent with the same column names and data types as spy.agent, but with zero rows. "
        "From the Spy ERD, spy.agent has: agent_id, first, middle, last, address, city, country, salary, clearance_id — "
        "spy.newagent gets those nine columns with matching types and no data."
    ),
    "WHERE TRUE = FALSE is always false, so no rows from agent are copied. I get the column structure only (CREATE TABLE ... AS SELECT with an empty result).": (
        "WHERE TRUE = FALSE is always false, so no rows from spy.agent are copied. "
        "I get the column structure only (CREATE TABLE ... AS SELECT with an empty result)."
    ),
    "[SCREENSHOT: \\d agent and \\d newagent in psql, or DBeaver column list for both tables — show the nine columns above match (agent_id, first, middle, last, address, city, country, salary, clearance_id)]": (
        "[SCREENSHOT: \\d spy.agent and \\d spy.newagent in psql, or DBeaver column list for spy.agent and spy.newagent — "
        "show the nine columns above match (agent_id, first, middle, last, address, city, country, salary, clearance_id)]"
    ),
    "No. CREATE TABLE ... AS SELECT copies columns and data types only; it does not copy constraints from agent. On the ERD, agent has a primary key on agent_id and a foreign key clearance_id -> securityclearance.sc_id — those are not recreated on newagent unless I add them with ALTER TABLE.": (
        "No. CREATE TABLE ... AS SELECT copies columns and data types only; it does not copy constraints from spy.agent. "
        "On the ERD, spy.agent has a primary key on agent_id and a foreign key clearance_id -> spy.securityclearance.sc_id — "
        "those are not recreated on spy.newagent unless I add them with ALTER TABLE."
    ),
    "9c. Include a screenshot showing that newagent has the same columns as agent.": (
        "9c. Include a screenshot showing that spy.newagent has the same columns as spy.agent."
    ),
}


def set_paragraph(paragraph, text, *, mono=False):
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = text
    run.font.name = "Consolas" if mono else "Calibri"
    run.font.size = Pt(9 if mono else 11)


def is_sql_line(line: str) -> bool:
    s = line.strip()
    if not s:
        return False
    if s.startswith("SQL") or s.startswith("Explanation") or s.startswith("Rows") or s.startswith("["):
        return False
    return (
        s.startswith("SELECT")
        or s.startswith("FROM")
        or s.startswith("WHERE")
        or s.startswith("INNER")
        or s.startswith("GROUP")
        or s.startswith("HAVING")
        or s.startswith("CREATE")
        or s.startswith("DROP")
        or s.startswith(");")
        or s.startswith("  ")
        or s == ");"
    )


def update_part2(doc: Document) -> None:
    paras = doc.paragraphs
    start = end = None
    for i, p in enumerate(paras):
        t = p.text.strip()
        if t == "Answer 6:":
            start = i
        if start is not None and t.startswith("Part 3:"):
            end = i
            break
    if start is None or end is None:
        raise RuntimeError("Could not find Answer 6 / Part 3 block")

    needed = len(SPY_ANSWER_LINES)
    available = end - start
    if needed > available:
        raise RuntimeError(f"Need {needed} paragraphs but only {available} slots")

    for j, line in enumerate(SPY_ANSWER_LINES):
        set_paragraph(paras[start + j], line, mono=is_sql_line(line))
    for j in range(start + needed, end):
        set_paragraph(paras[j], "")
    print(f"Part 2: updated paragraphs {start}-{start + needed - 1}")


def update_part3(doc: Document) -> None:
    q9_changed = 0
    for para in doc.paragraphs:
        raw = para.text
        new = raw
        for old, new_val in Q9_REPLACEMENTS.items():
            if old in new:
                new = new.replace(old, new_val)
        if new != raw:
            para.text = new
            q9_changed += 1
    print(f"Part 3: updated {q9_changed} paragraphs")


def main() -> None:
    doc = Document(str(DOC))
    update_part2(doc)
    update_part3(doc)
    doc.save(str(DOC))
    print(f"Saved {DOC.name}")


if __name__ == "__main__":
    main()
