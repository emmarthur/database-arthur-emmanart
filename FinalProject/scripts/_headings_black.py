# -*- coding: utf-8 -*-
"""Force Heading / Title runs to black (overrides theme blue)."""
from docx import Document
from docx.shared import RGBColor

BLACK = RGBColor(0x00, 0x00, 0x00)

TARGET_STYLES = {"Title", "Subtitle"}

from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parents[2]
FP = REPO_ROOT / "FinalProject"
PATHS = [
    REPO_ROOT / "Homework" / "Homework 3-emmanart.docx",
    FP / "deliverables" / "deliverable-3-emmanart.docx",
    FP / "deliverables" / "Deliverable-3-Instructions.docx",
    FP / "deliverables" / "Deliverable2_emmanart.docx",
]


def blacken_paragraph_runs(paragraph):
    for run in paragraph.runs:
        run.font.color.rgb = BLACK


def blacken_paragraphs(paragraphs):
    for p in paragraphs:
        name = p.style.name if p.style else ""
        if name.startswith("Heading") or name in TARGET_STYLES:
            blacken_paragraph_runs(p)


def blacken_doc(doc):
    blacken_paragraphs(doc.paragraphs)
    for section in doc.sections:
        for block in (section.header, section.footer):
            blacken_paragraphs(block.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                blacken_paragraphs(cell.paragraphs)


def main():
    for path in PATHS:
        try:
            doc = Document(path)
            blacken_doc(doc)
            doc.save(path)
            print("Updated:", path)
        except FileNotFoundError:
            print("Skip (missing):", path)


if __name__ == "__main__":
    main()
