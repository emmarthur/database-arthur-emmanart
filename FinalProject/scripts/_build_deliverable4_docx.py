# -*- coding: utf-8 -*-
"""Build Final Project Deliverable 4 write-up docx."""
import importlib.util
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parents[2]
FP = REPO_ROOT / "FinalProject"
OUT = FP / "deliverables" / "docx" / "Final_Project_Deliverable_4-WriteUp.docx"
MIG = FP / "soccer_data" / "migration_artifacts"
BUILD3 = FP / "scripts" / "_build_deliverable3_docx.py"

HEADING_BLACK = RGBColor(0, 0, 0)

VIEWS_SQL = """CREATE OR REPLACE VIEW soccer_proj.league_match_summary_view AS
SELECT
  l.id AS league_id,
  l.name AS league_name,
  COUNT(*) AS total_matches,
  AVG(m.home_team_goal)::numeric(10,2) AS avg_home_goals,
  AVG(m.away_team_goal)::numeric(10,2) AS avg_away_goals,
  AVG((m.home_team_goal + m.away_team_goal)::numeric(10,2)) AS avg_total_goals
FROM soccer_proj."Match_tbl" m
JOIN soccer_proj."League" l ON l.id = m.league_id
GROUP BY l.id, l.name;

CREATE OR REPLACE VIEW soccer_proj.match_context_view AS
SELECT
  m.id AS match_id,
  m.season,
  m.stage,
  m.date,
  c.name AS country_name,
  l.name AS league_name,
  ht.team_long_name AS home_team,
  at.team_long_name AS away_team,
  m.home_team_goal,
  m.away_team_goal
FROM soccer_proj."Match_tbl" m
JOIN soccer_proj."Country" c ON c.id = m.country_id
JOIN soccer_proj."League" l ON l.id = m.league_id
JOIN soccer_proj."Team" ht ON ht.team_api_id = m.home_team_api_id
JOIN soccer_proj."Team" at ON at.team_api_id = m.away_team_api_id;"""

INDEXES_SQL = """CREATE INDEX idx_match_league_id ON soccer_proj."Match_tbl" (league_id);
CREATE INDEX idx_match_season ON soccer_proj."Match_tbl" (season);
CREATE INDEX idx_match_home_team ON soccer_proj."Match_tbl" (home_team_api_id);
CREATE INDEX idx_match_away_team ON soccer_proj."Match_tbl" (away_team_api_id);
CREATE INDEX idx_league_country_id ON soccer_proj."League" (country_id);
CREATE INDEX idx_player_attr_player_api ON soccer_proj."Player_Attributes" (player_api_id);
CREATE INDEX idx_team_attr_team_api ON soccer_proj."Team_Attributes" (team_api_id);"""

FK_SAFE_SQL = """ALTER TABLE soccer_proj."League"
  ADD CONSTRAINT League_fk_country
  FOREIGN KEY (country_id) REFERENCES soccer_proj."Country" (id);

ALTER TABLE soccer_proj."Match_tbl"
  ADD CONSTRAINT Match_tbl_fk_league
  FOREIGN KEY (league_id) REFERENCES soccer_proj."League" (id);

ALTER TABLE soccer_proj."Match_tbl"
  ADD CONSTRAINT Match_tbl_fk_home_team
  FOREIGN KEY (home_team_api_id) REFERENCES soccer_proj."Team" (team_api_id);

ALTER TABLE soccer_proj."Match_tbl"
  ADD CONSTRAINT Match_tbl_fk_away_team
  FOREIGN KEY (away_team_api_id) REFERENCES soccer_proj."Team" (team_api_id);

ALTER TABLE soccer_proj."Team_Attributes"
  ADD CONSTRAINT Team_Attributes_fk_team_api
  FOREIGN KEY (team_api_id) REFERENCES soccer_proj."Team" (team_api_id);

ALTER TABLE soccer_proj."Player_Attributes"
  ADD CONSTRAINT Player_Attributes_fk_player_api
  FOREIGN KEY (player_api_id) REFERENCES soccer_proj."Player" (player_api_id);"""

ROW_COUNTS = [
    ("Country", 11),
    ("League", 11),
    ("Team", 299),
    ("Player", 11060),
    ("Team_Attributes", 1458),
    ("Player_Attributes", 183978),
    ("Match_tbl", 25979),
    ("league_match_summary_view", 11),
]

TOTAL_ROWS = sum(r for n, r in ROW_COUNTS if n != "league_match_summary_view")


def load_questions_sql():
    spec = importlib.util.spec_from_file_location("d3", BUILD3)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod.QUESTIONS_SQL


def add_title(doc, text, size=16):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_h(doc, text, level=1):
    doc.add_heading(text, level=level)
    p = doc.paragraphs[-1]
    for r in p.runs:
        r.font.name = "Calibri"
        r.font.color.rgb = HEADING_BLACK


def add_p(doc, text, italic=False):
    p = doc.add_paragraph(text)
    for r in p.runs:
        r.font.name = "Calibri"
        r.font.size = Pt(11)
        r.italic = italic


def add_sql(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(8)
    p.paragraph_format.left_indent = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def add_file_sql(doc, path, max_chars=None):
    text = path.read_text(encoding="utf-8")
    if max_chars and len(text) > max_chars:
        add_p(
            doc,
            f"(Full DDL is in {path.name}; excerpt below.)",
            italic=True,
        )
        text = text[:max_chars] + "\n\n-- ... remainder omitted in doc; see migration_artifacts folder ..."
    add_sql(doc, text)


def build():
    doc = Document()
    add_title(doc, "CS 486/586 Final Project\nDeliverable 4: Final Write-Up", 15)
    add_p(doc, "Student: Emmanuel Arthur / emmanart", italic=True)
    add_p(
        doc,
        "Video presentation link: [PASTE YOUR WORKING VIDEO URL HERE]",
        italic=True,
    )
    add_p(
        doc,
        "Data source: European Soccer Database (Kaggle) "
        "https://www.kaggle.com/datasets/hugomathien/soccer",
    )

    add_h(doc, "4.1 Final project write-up", 1)
    add_p(
        doc,
        "This document pulls together Deliverables 1 through 3 for my European soccer "
        "database in PostgreSQL schema soccer_proj.",
    )

    add_h(doc, "4.1.1 Final ER diagram (as implemented)", 2)
    add_p(
        doc,
        "I implemented seven base tables in soccer_proj: Country, League, Team, Player, "
        "Team_Attributes, Player_Attributes, and Match_tbl. Country links to League; League "
        "and Country link to Match_tbl; home and away teams link to Match_tbl through "
        "team_api_id; Player links to Match_tbl through 22 lineup columns; attribute tables "
        "link to Team and Player.",
    )
    add_p(
        doc,
        "Compared to my early conceptual ERD, the main changes were: (1) keeping the wide "
        "Match_tbl shape from the Kaggle SQLite export instead of splitting lineups into a "
        "separate table, (2) enforcing only PostgreSQL-valid foreign keys after migration "
        "(some SQLite metadata FKs were logical-only), and (3) adding views for reporting "
        "instead of extra physical tables.",
    )
    add_p(
        doc,
        "Paste your DBeaver ER diagram screenshot here, or export from SoccerProj-ERD.md.",
        italic=True,
    )
    add_p(doc, "Implemented relationships (enforced where noted):")
    add_p(doc, "• Country (1) to League (many) via League.country_id")
    add_p(doc, "• League (1) to Match_tbl (many) via Match_tbl.league_id")
    add_p(doc, "• Country (1) to Match_tbl (many) via Match_tbl.country_id")
    add_p(doc, "• Team (1) to Match_tbl (many) for home_team_api_id and away_team_api_id")
    add_p(doc, "• Team (1) to Team_Attributes (many) via team_api_id")
    add_p(doc, "• Player (1) to Player_Attributes (many) via player_api_id")
    add_p(doc, "• Player (1) to Match_tbl lineup columns (home_player_1..11, away_player_1..11)")

    add_h(doc, "4.1.2 CREATE TABLE, keys, and VIEW statements", 2)
    add_p(doc, "Schema and tables (from postgres_schema_tables_only.sql):")
    add_file_sql(doc, MIG / "postgres_schema_tables_only.sql")
    add_p(doc, "Foreign keys applied in PostgreSQL (safe subset):")
    add_sql(doc, FK_SAFE_SQL)
    add_p(
        doc,
        "Additional Match_tbl player FKs exist in postgres_schema.sql when parent "
        "player_api_id values are unique-enough for PostgreSQL enforcement.",
    )
    add_p(doc, "Views:")
    add_sql(doc, VIEWS_SQL)

    add_h(doc, "4.1.3 CREATE INDEX statements (graduate)", 2)
    add_p(
        doc,
        "I added indexes on common join and filter columns used in my Deliverable 3 queries.",
    )
    add_sql(doc, INDEXES_SQL)

    add_h(doc, "4.1.4 How I populated the database", 2)
    add_p(
        doc,
        "I started from the Kaggle European Soccer SQLite file (database.sqlite). I used "
        "sqlite_to_postgres_prep.py in the FinalProject folder to export CSV files and "
        "generate PostgreSQL DDL under soccer_data/migration_artifacts/.",
    )
    add_p(doc, "Load steps in DBeaver / PostgreSQL:")
    add_p(doc, "1. drop schema if exists soccer_proj cascade; create schema soccer_proj;")
    add_p(doc, "2. Run postgres_schema_tables_only.sql to create empty tables.")
    add_p(doc, "3. Import each CSV from migration_artifacts/csv/ into the matching table.")
    add_p(doc, "4. Run postgres_schema_fks_safe.sql (or safe rerunnable version) for FKs.")
    add_p(doc, "5. Create views league_match_summary_view and match_context_view.")
    add_p(doc, "6. Create indexes listed in section 4.1.3.")
    add_p(
        doc,
        "I did not mock the data. Row counts match migration_summary.txt and Deliverable 2.",
    )

    add_h(doc, "4.1.5 English questions, SQL, and results", 2)
    add_p(
        doc,
        "Below are my 30 questions from Deliverable 1.2, the SQL I used in Deliverable 3, "
        "and space for query output. For each query I ran it in DBeaver and pasted the "
        "first 10 rows (or the full result when smaller).",
    )
    add_h(doc, "Question revisions", 3)
    add_p(
        doc,
        "I did not replace any question entirely. I adjusted how a few map to the loaded data:",
    )
    add_p(
        doc,
        "• Q25: I used preferred_foot from Player_Attributes instead of parsing lineup "
        "coordinate columns.",
    )
    add_p(
        doc,
        "• Q29: possession is stored as text/XML, so I used completeness-style logic "
        "instead of parsing possession share.",
    )
    add_p(
        doc,
        "• Q10, Q11, Q24, Q28: I wrote full window-function / UNION ALL SQL in Deliverable 3.",
    )

    questions = load_questions_sql()
    for tag, english, sql in questions:
        add_h(doc, f"{tag}. {english}", 3)
        add_p(doc, "SQL:")
        add_sql(doc, sql)
        add_p(
            doc,
            "Result (first 10 rows): [paste DBeaver screenshot or table output here]",
            italic=True,
        )

    add_h(doc, "4.1.6 Five rows per table and row counts", 2)
    add_p(doc, "Row count per table/view:")
    add_sql(
        doc,
        """SELECT 'Country' AS object_name, COUNT(*) AS row_count FROM soccer_proj."Country"
UNION ALL SELECT 'League', COUNT(*) FROM soccer_proj."League"
UNION ALL SELECT 'Team', COUNT(*) FROM soccer_proj."Team"
UNION ALL SELECT 'Player', COUNT(*) FROM soccer_proj."Player"
UNION ALL SELECT 'Team_Attributes', COUNT(*) FROM soccer_proj."Team_Attributes"
UNION ALL SELECT 'Player_Attributes', COUNT(*) FROM soccer_proj."Player_Attributes"
UNION ALL SELECT 'Match_tbl', COUNT(*) FROM soccer_proj."Match_tbl"
UNION ALL SELECT 'league_match_summary_view', COUNT(*) FROM soccer_proj.league_match_summary_view
ORDER BY object_name;""",
    )
    add_p(doc, "Expected counts from my loaded database:")
    for name, cnt in ROW_COUNTS:
        add_p(doc, f"• {name}: {cnt:,}")
    add_p(doc, "Sample: five rows from each base table:")
    tables = [
        "Country",
        "League",
        "Team",
        "Player",
        "Team_Attributes",
        "Player_Attributes",
        "Match_tbl",
    ]
    for t in tables:
        add_p(doc, f"{t}:")
        add_sql(doc, f'SELECT * FROM soccer_proj."{t}" LIMIT 5;')
        add_p(doc, "[paste 5-row screenshot here]", italic=True)

    add_h(doc, "4.1.7 Total row count (all tables)", 2)
    add_p(
        doc,
        f"Sum of rows across the seven base tables: {TOTAL_ROWS:,} "
        "(excludes the view, which is derived from Match_tbl and League).",
    )
    add_sql(
        doc,
        """SELECT SUM(row_count) AS total_rows_all_tables
FROM (
  SELECT COUNT(*) AS row_count FROM soccer_proj."Country"
  UNION ALL SELECT COUNT(*) FROM soccer_proj."League"
  UNION ALL SELECT COUNT(*) FROM soccer_proj."Team"
  UNION ALL SELECT COUNT(*) FROM soccer_proj."Player"
  UNION ALL SELECT COUNT(*) FROM soccer_proj."Team_Attributes"
  UNION ALL SELECT COUNT(*) FROM soccer_proj."Player_Attributes"
  UNION ALL SELECT COUNT(*) FROM soccer_proj."Match_tbl"
) counts;""",
    )
    add_p(doc, "[paste screenshot of total_rows_all_tables query here]", italic=True)

    add_h(doc, "Final presentation checklist (video)", 2)
    add_p(doc, "• Dataset overview and why I chose European soccer")
    add_p(doc, "• Initial ERD vs final ERD (what changed)")
    add_p(doc, "• Demo tables + row counts")
    add_p(doc, "• Demo one intermediate query and one advanced query")
    add_p(doc, "• Hardest part, what I learned, SQL confidence, next steps with more time")

    apply_black_headings(doc)
    doc.save(OUT)
    print("Wrote", OUT)


def apply_black_headings(doc):
    for p in doc.paragraphs:
        st = p.style.name if p.style else ""
        if st.startswith("Heading") or st in ("Title", "Subtitle"):
            for r in p.runs:
                r.font.color.rgb = HEADING_BLACK


if __name__ == "__main__":
    build()
