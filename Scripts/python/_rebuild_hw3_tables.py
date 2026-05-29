# -*- coding: utf-8 -*-
"""Replace shorthand schema lines with reconstructed Word tables for HW3."""
from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.text.paragraph import Paragraph


def delete_paragraph(paragraph):
    el = paragraph._element
    el.getparent().remove(el)


def insert_paragraph_after(paragraph, text="", bold=False):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        r = new_para.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(11)
        r.bold = bold
    return new_para


def insert_paragraph_after_tbl(table, doc, text="", bold=False):
    """Insert a paragraph immediately after a Table object."""
    new_p = OxmlElement("w:p")
    table._tbl.addnext(new_p)
    new_para = Paragraph(new_p, doc.element.body)
    if text:
        r = new_para.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(11)
        r.bold = bold
    return new_para


def move_table_after(table, paragraph):
    paragraph._p.addnext(table._tbl)


def fill_header_row(table, headers):
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for run in p.runs:
                run.bold = True


def build_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    fill_header_row(t, headers)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
    return t


def delete_block_after_title(doc, title_strip: str, stop_prefix: str):
    paras = list(doc.paragraphs)
    recording = False
    to_delete = []
    for p in paras:
        t = p.text.strip()
        if t == title_strip:
            recording = True
        if recording:
            if t.startswith(stop_prefix) and title_strip not in t:
                break
            to_delete.append(p)
    for p in reversed(to_delete):
        delete_paragraph(p)


def find_paragraph(doc, prefix: str):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    return None


def main():
    from pathlib import Path

    path = Path(__file__).resolve().parents[2] / "Homework" / "docx" / "Homework 3-emmanart.docx"
    doc = Document(path)

    # --- Remove old shorthand blocks (keep intro sentences before them)
    delete_block_after_title(doc, "Tables (2NF):", "Part 3")
    delete_block_after_title(doc, "Tables (3NF):", "SQL for DBeaver")

    # --- 1NF: caption before existing 1NF table (between PK line and Part 2)
    pk_p = find_paragraph(doc, "Primary key: (StudentID, CourseID, Semester).")
    if pk_p and not find_paragraph(doc, "Reconstructed 1NF registration table"):
        insert_paragraph_after(
            pk_p,
            "Reconstructed 1NF registration table (same scenario data, just exploded so nothing’s multi-valued):",
        )

    # --- 2NF: tables after intro paragraph
    p2 = find_paragraph(doc, "Pulled partial deps off the big enrollment key")
    if not p2:
        raise SystemExit("Could not find 2NF intro paragraph")

    cur = insert_paragraph_after(
        p2,
        "Below are the actual split tables with the sample rows filled in — way easier to read than one long schema line.",
    )
    cur = insert_paragraph_after(cur, "Student — primary key = StudentID")
    t_stu = build_table(doc, ["StudentID", "StudentName", "Major"], student_rows)
    move_table_after(t_stu, cur)
    cur = insert_paragraph_after_tbl(t_stu, doc, "Course — primary key = CourseID")

    t_co = build_table(doc, ["CourseID", "CourseName"], course_rows)
    move_table_after(t_co, cur)
    cur = insert_paragraph_after_tbl(
        t_co,
        doc,
        "CourseOffering — primary key = (CourseID, Semester); ties instructor to a specific offering.",
    )

    t_off = build_table(
        doc,
        ["CourseID", "Semester", "InstructorName"],
        offering_2nf_rows,
    )
    move_table_after(t_off, cur)
    cur = insert_paragraph_after_tbl(
        t_off,
        doc,
        "Enrollment — primary key = (StudentID, CourseID, Semester); grade lives here only.",
    )

    t_en = build_table(
        doc,
        ["StudentID", "CourseID", "Semester", "Grade"],
        enrollment_rows,
    )
    move_table_after(t_en, cur)
    cur = insert_paragraph_after_tbl(
        t_en,
        doc,
        "Foreign keys: Enrollment.StudentID → Student; Enrollment.CourseID → Course; Enrollment lines up with CourseOffering through (CourseID, Semester).",
    )

    # --- 3NF tables
    p3 = find_paragraph(doc, "Split instructor out so we")
    if not p3:
        raise SystemExit("Could not find 3NF intro paragraph")

    cur = insert_paragraph_after(
        p3,
        "Same story but Instructor is its own table so we’re not treating instructor text as random strings on the offering.",
    )
    cur = insert_paragraph_after(cur, "Student")
    t_s3 = build_table(doc, ["StudentID", "StudentName", "Major"], student_rows)
    move_table_after(t_s3, cur)
    cur = insert_paragraph_after_tbl(t_s3, doc, "Course")

    t_c3 = build_table(doc, ["CourseID", "CourseName"], course_rows)
    move_table_after(t_c3, cur)
    cur = insert_paragraph_after_tbl(t_c3, doc, "Instructor — primary key = InstructorID")

    t_ins = build_table(doc, ["InstructorID", "InstructorName"], instructor_rows)
    move_table_after(t_ins, cur)
    cur = insert_paragraph_after_tbl(
        t_ins,
        doc,
        "CourseOffering — primary key = (CourseID, Semester); InstructorID → Instructor.",
    )

    t_of3 = build_table(
        doc,
        ["CourseID", "Semester", "InstructorID"],
        offering_3nf_rows,
    )
    move_table_after(t_of3, cur)
    cur = insert_paragraph_after_tbl(t_of3, doc, "Enrollment")

    t_e3 = build_table(
        doc,
        ["StudentID", "CourseID", "Semester", "Grade"],
        enrollment_rows,
    )
    move_table_after(t_e3, cur)
    insert_paragraph_after_tbl(
        t_e3,
        doc,
        "FKs in plain English: enrollment rows point at student + offering; offering rows point at course + instructor; instructor row holds the name once.",
    )

    doc.save(path)
    print("Updated:", path)


# Sample data (matches homework scenario)
student_rows = [
    (101, "Maya Patel", "CS"),
    (102, "Daniel Wong", "DS"),
    (103, "Taylor Alison", "CS"),
]

course_rows = [
    ("CS101", "Intro to CS"),
    ("CS205", "DBMS"),
    ("MATH201", "Discrete Math"),
]

offering_2nf_rows = [
    ("CS101", "Fall 2025", "Lee"),
    ("CS205", "Fall 2025", "Smith"),
    ("MATH201", "Fall 2025", "Brown"),
]

offering_3nf_rows = [
    ("CS101", "Fall 2025", 1),
    ("CS205", "Fall 2025", 2),
    ("MATH201", "Fall 2025", 3),
]

instructor_rows = [
    (1, "Lee"),
    (2, "Smith"),
    (3, "Brown"),
]

enrollment_rows = [
    (101, "CS101", "Fall 2025", "A"),
    (101, "CS205", "Fall 2025", "B+"),
    (102, "CS205", "Fall 2025", "A-"),
    (103, "CS101", "Fall 2025", "B"),
    (103, "MATH201", "Fall 2025", "A"),
]


if __name__ == "__main__":
    main()
